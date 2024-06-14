from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from openai import OpenAI


class Author:
    def __init__(self, api_key):
        self.client = OpenAI(api_key=api_key)

    def generate_section(self, prompt, context):
        completion = self.client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": """
                You are a helpful assistant that will help us draft a review article.

                Remember: 
                - The abstract should summarize the whole paper. 
                - The introduction should clearly introduce the content and lay the direction of this paper.
                - The review is where you discuss all previous work that has been done. This should be the largest part of the paper. Split this into different subtopics. 
                - The conclusion is where you summarize and conclude the paper.
                - The sources should include a list of all relevant sources, organized in MLA format.
                - Do not include section names in the content
                """},
                {"role": "user", "content": context},
                {"role": "user", "content": prompt}
            ]
        )
        return completion.choices[0].message.content

    def generate_outline(self, topic, description):
        prompt = f"Create an outline for a review article on the topic '{topic}' with the following detailed description: {description}. The outline should include the sections: Title, Abstract, Introduction, Body, Conclusion, and Sources."
        outline = self.generate_section(prompt, "")
        return outline

    def draft_paper(self, outline):
        sections = ["Title", "Abstract", "Introduction", "Review", "Conclusion", "Sources"]
        paper_dict = {"Outline": outline}
        context = outline

        for section in sections:
            prompt = f"Draft the {section} section of the paper based on the outline."
            section_content = self.generate_section(prompt, context)
            paper_dict[section] = section_content
            context += f"\n\n{section}:\n{section_content}"

        return paper_dict

    def save_to_word(self, content, filename="Drafted_Paper.docx"):
        doc = Document()

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        title = content.pop('Title', '<Insert Title Here>')
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(14)
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        content.pop('Outline', None)

        for section, section_content in content.items():
            section_heading = doc.add_heading(level=1)
            heading_run = section_heading.add_run(section)
            heading_run.bold = True
            heading_run.font.name = 'Times New Roman'
            heading_run.font.size = Pt(12)

            for line in section_content.split("\n"):
                if line.strip() == "":
                    doc.add_paragraph()
                else:
                    paragraph = doc.add_paragraph(line)
                    paragraph.style = doc.styles['Normal']
                    paragraph_font = paragraph.runs[0].font
                    paragraph_font.size = Pt(10)

        doc.save(filename)
